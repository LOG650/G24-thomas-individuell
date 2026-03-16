"""
Bygg Feynman-hjelpedokument (Word) for LOG650-rapporten.
Forklarer hele rapporten, alle formler, figurer og tabeller
med Feynman-teknikk: enkelt språk, analogier, trinn for trinn.

Kjør: cd "005 report" && py build_feynman_guide.py
"""

import os, sys
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# ─── Farger ───
TITLE_COLOR = RGBColor(0x1A, 0x2A, 0x44)
GREEN = RGBColor(0x1E, 0x7D, 0x45)
BLUE = RGBColor(0x0B, 0x3D, 0x8C)
ORANGE = RGBColor(0xD6, 0x89, 0x10)
RED = RGBColor(0xB0, 0x3A, 0x2E)
GREY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0, 0, 0)

OUTPUT = 'Feynman_Hjelpedokument_LOG650.docx'

doc = Document()

# ─── Standardstiler ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = TITLE_COLOR
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, bold=False, italic=False, color=None, indent_cm=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p


def add_formula_box(label, formula, explanation):
    """Legger til en formel i en visuell boks med forklaring."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'{label}')
    r.bold = True
    r.font.color.rgb = BLUE
    r.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.5)
    r2 = p2.add_run(formula)
    r2.font.name = 'Consolas'
    r2.font.size = Pt(11)
    r2.font.color.rgb = BLACK

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(1.5)
    p3.paragraph_format.space_after = Pt(8)
    r3 = p3.add_run(explanation)
    r3.font.size = Pt(10)
    r3.font.color.rgb = GREY
    r3.italic = True


def add_feynman_tip(text):
    """Legger til en Feynman-analogi/tips-boks."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    icon = p.add_run('Tenk det slik: ')
    icon.bold = True
    icon.font.color.rgb = GREEN
    icon.font.size = Pt(11)
    tip = p.add_run(text)
    tip.font.size = Pt(11)
    tip.italic = True
    tip.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table_simple(headers, rows):
    """Enkel tabell med header og rader."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing


# ══════════════════════════════════════════════════════════════════
# INNHOLD
# ══════════════════════════════════════════════════════════════════

# ─── FORSIDE ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('Feynman-hjelpedokument')
r.font.size = Pt(28)
r.font.color.rgb = TITLE_COLOR
r.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('LOG650 Forskningsprosjekt')
r2.font.size = Pt(16)
r2.font.color.rgb = BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(20)
r3 = p3.add_run(
    'Fra lokalt forsyningslager til regional sentralforsyning:\n'
    'Multikriterieklassifisering og klyngeanalyse for\n'
    'identifisering av overføringskandidater ved Helse Bergen'
)
r3.font.size = Pt(13)
r3.italic = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(30)
r4 = p4.add_run(
    'Dette dokumentet forklarer rapporten, alle formler, figurer og tabeller\n'
    'med Feynman-teknikken: som om du forklarer det til noen uten forkunnskaper.'
)
r4.font.size = Pt(11)
r4.font.color.rgb = GREY

doc.add_page_break()

# ─── INNHOLDSFORTEGNELSE ───
add_heading('Innhold', 1)
toc_items = [
    'Del 1 – Hva handler rapporten om? (Helheten)',
    'Del 2 – Hele rapporten, del for del',
    '    Forord og sammendrag',
    '    Kapittel 1–9',
    '    Vedlegg A–C',
    'Del 3 – Alle formler forklart',
    '    3.1 ABC-analyse (verdiklassifisering)',
    '    3.2 XYZ-klassifisering (variasjonskoeffisient)',
    '    3.3 EOQ – Wilson-modellen (optimal bestillingsmengde)',
    '    3.4 Avviksformel (hvor mye bestiller vi feil?)',
    '    3.5 Totalkostnadsfunksjon',
    '    3.6 Besparelsesformel',
    '    3.7 K-means featurevektor',
    '    3.8 Silhouette-score',
    'Del 4 – Alle figurer forklart (Figur 1–12)',
    'Del 5 – Alle tabeller forklart (Tabell 1–15)',
    'Del 6 – Ordliste og nøkkelbegreper',
]
for item in toc_items:
    p = doc.add_paragraph()
    indent = 1.5 if item.startswith('    ') else 0.5
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item.strip())
    r.font.size = Pt(11)
    if not item.startswith('    '):
        r.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DEL 1 – HELHETEN
# ══════════════════════════════════════════════════════════════════
add_heading('Del 1 – Hva handler rapporten om?', 1)

add_para(
    'Feynman-teknikken sier: Hvis du ikke kan forklare noe enkelt, forstår du '
    'det ikke godt nok. Denne delen forklarer hele rapporten som om du aldri '
    'har hørt om lagerstyring, SAP eller statistikk.'
)

add_heading('Den korte versjonen', 2)
add_para(
    'Helse Bergen (Haukeland sykehus) har et lager med over 700 ulike varer – '
    'hansker, sprøyter, bandasjer og annet medisinsk forbruksmateriell. '
    'Akkurat nå bestiller de alt lokalt, men det finnes et nytt regionalt '
    'sentrallager (HVFS) som kan ta over noen av varene. Spørsmålet er: '
    'HVILKE varer bør flyttes til sentrallageret, og HVOR MYE penger kan '
    'sykehuset spare på det?'
)

add_feynman_tip(
    'Tenk på det som en familie som handler dagligvarer. Noen ting kjøper '
    'du hver uke i nærbutikken (melk, brød). Andre ting er billigere i '
    'storpakninger fra Costco (dopapir, vaskemiddel). Rapporten prøver å '
    'finne ut hvilke «sykehus-varer» som bør kjøpes lokalt og hvilke som '
    'bør komme fra «Costco-lageret» (HVFS).'
)

add_heading('Hva gjør rapporten konkret?', 2)
add_para(
    'Rapporten bruker fire analysemetoder – tenk på dem som fire ulike '
    '«tester» hver vare må gjennom:'
)

items = [
    ('1. ABC-analyse (Verditest)', 'Sorterer varene etter hvor mye penger de binder opp. '
     'A-varer = de dyreste 80 % av verdien. C-varer = de billigste.'),
    ('2. XYZ-klassifisering (Forutsigbarhetstest)', 'Måler hvor jevnt forbruket er. '
     'X = stabilt forbruk (f.eks. hansker brukes hver dag). Z = uforutsigbart.'),
    ('3. EOQ-avviksanalyse (Bestillingstest)', 'Sjekker om sykehuset bestiller for '
     'ofte eller for sjelden. Hvis du bestiller 50 ganger i året men optimalt '
     'er 10 ganger, kaster du bort penger på ordrekostnader.'),
    ('4. K-means klyngeanalyse (Mønstergjenkjenning)', 'En maskinlæringsalgoritme '
     'som grupperer varer som «ligner» hverandre basert på verdi, stabilitet '
     'og bestillingsavvik.'),
]
for title, desc in items:
    add_para(title, bold=True)
    add_para(desc, indent_cm=1)

add_para(
    'Til slutt kombineres alle fire testene i en regelmotor – et sett med '
    '8 regler som gir hver vare én av fire anbefalinger:'
)

recs = [
    ('OVERFØR TIL HVFS', '145 artikler', 'Flytt til sentrallageret'),
    ('BEHOLD LOKALT', '257 artikler', 'Behold på sykehuset'),
    ('VURDER NÆRMERE', '284 artikler', 'Krever manuell vurdering'),
    ('MANGLER DATA', '23 artikler', 'For lite data til å avgjøre'),
]
add_table_simple(
    ['Anbefaling', 'Antall', 'Hva betyr det?'],
    recs
)

add_para(
    'Estimert besparelse: ca. 450 000 kr/år i base case – det sykehuset '
    'kan spare ved å gå fra «mange små bestillinger» til optimal '
    'bestillingspraksis gjennom sentralisering.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DEL 2 – KAPITTEL FOR KAPITTEL
# ══════════════════════════════════════════════════════════════════
add_heading('Del 2 – Hele rapporten, del for del', 1)

# ─── Forord og sammendrag ───
add_heading('Forord', 2)
add_para(
    'Forordet er forfatterens personlige innledning. Thomas Ekrem Jensen '
    'jobber til daglig som SAP MM-konsulent i Helse Vest IKT på LIBRA-prosjektet. '
    'Motivasjonen for oppgaven er gapet mellom data som allerede finnes i SAP '
    'og den innsikten som faktisk brukes til å ta beslutninger. Han takker '
    'veilederen, Helse Bergen for datatilgang, og samboeren for tålmodighet.'
)
add_feynman_tip(
    'Forordet svarer på «Hvem er denne personen, og hvorfor bryr han seg?» '
    'Svaret: Han ser daglig at SAP er fullt av data som ingen utnytter. '
    'Oppgaven er hans forsøk på å bevise at dataene faktisk kan brukes.'
)

add_heading('Sammendrag', 2)
add_para(
    'Sammendraget er hele rapporten komprimert til én side. Det inneholder:',
    bold=True
)
summary_points = [
    'Bakgrunn: LIBRA-prosjektet og HVFS gjør sentralisering aktuelt.',
    'Data: 709 aktive artikler fra 14 SAP-tabeller (2024–2025). '
    '297 inaktive artikler er ekskludert.',
    'Metode: Fire analyser (ABC, XYZ, EOQ, K-means) + regelmotor med 8 regler.',
    'Hovedresultat: 145 artikler anbefales overført, 257 beholdes lokalt, '
    '284 krever manuell vurdering, 23 mangler data.',
    'Besparelse: kr 451 515/år i base case (g = 75 %), basert på 117 artikler '
    'med dokumentert overbestilling. Intervall: kr 301 010 – kr 602 020.',
    'Konklusjon: Multikriterieklassifisering av SAP-data gir et operasjonelt '
    'anvendbart beslutningsgrunnlag for HVFS-overføring.',
]
for pt in summary_points:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(pt)
    r.font.size = Pt(11)

add_feynman_tip(
    'Sammendraget er «traileren» til rapporten. Hvis du bare leser én side, '
    'les denne – den gir alle nøkkeltallene og konklusjonen uten å måtte '
    'lese de 40 andre sidene.'
)

# ─── Kapitler ───
add_heading('Kapittel 1–9', 2)

chapters = [
    ('Kapittel 1 – Innledning', [
        'Bakgrunn: Sykehus bruker 30–40 % av driftskostnadene på logistikk. '
        'Mye penger kan spares ved å bestille smartere.',
        'Problemstilling: «Hvilke varer bør flyttes til HVFS, og hvor mye '
        'kan vi spare?»',
        'Avgrensninger: Kun LGORT 3001 ved Helse Bergen. Kun medisinsk '
        'forbruksmateriell (ikke medisin eller implantater). Data fra 2024–2025.',
        'Antagelser: Ordrekostnad S = 750 kr per bestilling. Lagerholdekostnad '
        '= 20 % av varepris per år. Disse er hentet fra litteraturen, ikke '
        'målt hos Helse Bergen.',
    ]),
    ('Kapittel 2 – Litteratur og teori', [
        'Gjennomgår 23 akademiske kilder som handler om lagerstyring i sykehus.',
        'Forklarer ABC (verdiklassifisering), XYZ (forbruksstabilitet), '
        'EOQ (optimal bestillingsmengde) og K-means (maskinlæring).',
        'Identifiserer et «gap» i litteraturen: Ingen har tidligere kombinert '
        'alle fire metodene på SAP-data i en sykehusstudie.',
        'Presenterer et konseptuelt rammeverk (Figur 1) som viser hvordan '
        'de fire metodene henger sammen.',
    ]),
    ('Kapittel 3 – Casebeskrivelse', [
        'Helse Bergen: Haukeland sykehus, det største i Helse Vest-regionen.',
        'HVFS: Et nytt regionalt sentrallager drevet av NorEngros. '
        'Planlegger avdelingspakkede leveranser (APL) – dvs. ferdigpakkede '
        'kasser direkte til sykepleieren.',
        'LIBRA-prosjektet: Helse Vest sin SAP-implementering som gir felles '
        'dataplattform på tvers av sykehusene.',
        'Problemet i dag: Ingen systematisk metode for å bestemme hvilke '
        'varer som bør sentraliseres. Beslutninger tas på «magefølelse».',
    ]),
    ('Kapittel 4 – Metode og data', [
        'Forskningsdesign: Kvantitativ casestudie – tall fra SAP-systemet, '
        'ingen intervjuer eller spørreundersøkelser.',
        'Data: 14 SAP-tabeller hentet via SE16H (en lesetransaksjon i SAP). '
        '1006 artikler i utgangspunktet, redusert til 709 aktive.',
        '8 datavalgsbeslutninger (D-01 til D-08) dokumenterer alle valg '
        'analytikeren tok under databehandlingen.',
        'KI-bruk: Claude ble brukt til kodestøtte, figurer og tekststruktur, '
        'men IKKE til å produsere tallene eller trekke faglige konklusjoner.',
    ]),
    ('Kapittel 5 – Modellering', [
        'Her beskrives alle formler og modeller matematisk. '
        '(Se Del 3 nedenfor for Feynman-forklaring av hver formel.)',
        'ABC-modell: Sorter etter verdi, del i tre klasser.',
        'XYZ-modell: Beregn CV (variasjonskoeffisient) fra månedlig forbruk.',
        'EOQ-modell: Wilson-formelen for optimal bestillingsmengde.',
        'K-means-modell: Tre features, log-transformert og z-score-normalisert.',
        'Regelmotor: 8 regler i prioritert rekkefølge.',
    ]),
    ('Kapittel 6 – Analyse', [
        'Her kjøres modellene på de 709 artiklene. Alt implementert i Python.',
        'ABC: 182 A-artikler, 184 B, 338 C. Pareto-mønster bekreftet.',
        'XYZ: 350 X (stabilt), 193 Y, 144 Z. Over halvparten er stabile.',
        'EOQ: 73 % av artiklene bestiller for ofte (FOR_MANGE_ORDRER).',
        'K-means: K = 3 klynger. Silhouette-score 0.383 (akseptabelt).',
        'Regelmotor: Kjørt sekvensielt. 145 anbefalt overført.',
    ]),
    ('Kapittel 7 – Resultater', [
        'Presenterer alle tallresultater i tabeller.',
        'ABC-fordeling (Tabell 8), XYZ-fordeling (Tabell 9).',
        'ZZXYZ-validering: Kun 33 % samsvar med SAP-systemet.',
        'EOQ: Samlet teoretisk kostnadsavvik kr 2 333 441/år for alle 487 artikler '
        '– dette er IKKE direkte realiserbar besparelse, men et mål på hvor mye '
        'bestillingspraksisen totalt sett avviker fra EOQ-optimalt nivå.',
        'K-means: Tre klynger med tydelige profiler.',
        'Regelmotor: 145 OVERFØR, 257 BEHOLD, 284 VURDER, 23 MANGLER DATA.',
        'Faktisk besparelsesestimat: Gjelder kun de 117 artiklene som er '
        'OVERFØR_HVFS OG har FOR_MANGE_ORDRER. Base case kr 451 515/år (g=75 %). '
        'Worst kr 301 010, Best kr 602 020. Merk: de resterende 28 OVERFØR-artikler '
        'inngår ikke i besparelsesberegningen.',
    ]),
    ('Kapittel 8 – Diskusjon', [
        'Sammenligner funnene med litteraturen – alle 6 nøkkelfunn stemmer '
        '(Pareto-mønster, XYZ-samsvar, silhouette-nivå, besparelsesnivå, osv.).',
        'Metodekritikk – datakvalitet: SAP-data er automatisk registrert og '
        'ikke utsatt for hukommelsesfeil, men forbruk som skjer utenom SAP '
        '(f.eks. nøduttak som ikke dokumenteres) fanges ikke opp.',
        'Metodekritikk – ABC-verdi: 204 artikler mangler EKPO-innkjøpsdata '
        'og bruker beregnet verdi (D_ANNUAL x UNIT_PRICE). Hvis standardprisen '
        'avviker fra faktisk innkjøpspris, blir ABC-rangeringen upresis.',
        'Manglende ekstern validering: Regelmotoranbefalingene er IKKE validert '
        'mot innkjøpsfaglig skjønn eller historiske overføringsbeslutninger. '
        'Modellen bør betraktes som strukturert beslutningsunderlag, ikke '
        'som en autorisert beslutning. En oppfølgingsfase med faglig '
        'gjennomgang er nødvendig.',
        'Begrenset generaliserbarhet: Resultatene gjelder kun WERKS 3300, '
        'LGORT 3001. Metoderammeverket er reproduserbart, men andre foretak '
        '(Stavanger, Fonna, Førde) har ulik sortimentssammensetning, '
        'leverandøravtaler og driftsmodeller – tallene kan ikke overføres direkte.',
        'Praktisk betydning: 145 artikler kan flyttes, men gevinstrealisering '
        'forutsetter at SAP MM-parametere (MRP-type, ordrekvantumsinnstillinger) '
        'faktisk justeres. Uten det fortsetter systemet med eksisterende frekvens.',
        'Svakheter: Realiseringsgrad g er et scenario, ikke en prognose. '
        'Leveringstidsdata mangler for 94 % av artiklene. '
        'COVID-19-ettervirkninger kan ha påvirket forbruksmønstre i 2024–2025.',
    ]),
    ('Kapittel 9 – Konklusjon', [
        'Svaret på problemstillingen: Ja, multikriterieklassifisering og '
        'klyngeanalyse fungerer for å identifisere HVFS-kandidater.',
        '4 anbefalinger: (1) Pilotoverføring, (2) Manuell gjennomgang av '
        'VURDER-gruppen, (3) Oppdater SAP-parametere, (4) Evaluer etter 12 mnd.',
        '3 forslag til videre forskning: ROP-modul, leverandørkonsolidering, '
        'replikering til andre sykehus i Helse Vest.',
    ]),
]

for ch_title, points in chapters:
    add_heading(ch_title, 2)
    for pt in points:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(pt)
        r.font.size = Pt(11)

# ─── Vedlegg ───
add_heading('Vedlegg A–C', 2)

add_heading('Vedlegg A – SAP-dataspesifikasjon', 3)
add_para(
    'En detaljert tabell over alle 14 SAP-tabeller med nøkkelfelter. '
    'Eksempel: MARD inneholder lagerbeholdning per lagersted (felter MATNR, '
    'WERKS, LGORT, LABST). MSEG inneholder varebevegelser (MATNR, MENGE, BWART).'
)
add_feynman_tip(
    'Vedlegg A er «oppskriften» for å hente de samme dataene om igjen. '
    'Hvis noen vil gjenta analysen om ett år med ferske data, starter de her.'
)

add_heading('Vedlegg B – Python-analyseverktøy', 3)
add_para(
    'Beskriver analyseverktøyet: hovedscriptet LOG650_analyse_v2_7.py '
    '(deterministisk pipeline) og 12 figurscripts (plot_*.py). '
    'Nøkkelbiblioteker: pandas, scikit-learn, matplotlib, openpyxl. '
    'Alt bruker random_state = 42 for reproduserbarhet – gjentatt kjøring '
    'med samme input gir identiske resultater.'
)
add_feynman_tip(
    'Reproduserbarhet betyr at hvem som helst kan kjøre det samme scriptet '
    'og få nøyaktig de samme tallene. Det er som en matoppskrift der '
    'ingrediensene og mengdene er presist angitt – ingen «tilsett etter smak».'
)

add_heading('Vedlegg C – Erklæring om bruk av kunstig intelligens', 3)
add_para(
    'Dokumenterer bruken av Claude (Anthropic) i tre deler av arbeidet: '
    '(1) kodestøtte og feilsøking i Python, (2) generering av figurer '
    '(layout, farger, aksetitler), og (3) strukturering av rapportteksten. '
    'Alle numeriske resultater, tabellverdier og analytiske konklusjoner '
    'er produsert av Python-scriptet fra SAP-kildedata, uavhengig av '
    'KI-verktøyet. Claude er ikke brukt som fagkilde.'
)
add_feynman_tip(
    'Vedlegg C er rapportens «ærlighetsattest» for KI-bruk. Det er påkrevd '
    'av Høgskolen i Molde og viser at forfatteren har kontroll på hva som '
    'er hans egne analyser og hva KI har bidratt med.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DEL 3 – ALLE FORMLER FORKLART
# ══════════════════════════════════════════════════════════════════
add_heading('Del 3 – Alle formler forklart', 1)

add_para(
    'Denne delen tar hver eneste formel i rapporten og forklarer den med '
    'Feynman-teknikk: hva den betyr i vanlige ord, hva hvert symbol er, '
    'og et konkret talleksempel.'
)

# ─── 3.1 ABC ───
add_heading('3.1 ABC-analyse (verdiklassifisering)', 2)

add_heading('Formel: Årsverdi per artikkel', 3)
add_formula_box(
    'Formel:',
    'v_i = D_i  x  UNIT_PRICE_i',
    'Årsverdien for en vare = hvor mange enheter vi bruker per år, ganget med prisen per enhet.'
)
add_para('Hva betyr symbolene?')
add_table_simple(
    ['Symbol', 'Betydning', 'Eksempel'],
    [
        ['v_i', 'Årsverdi for artikkel i (kroner)', 'kr 50 000'],
        ['D_i', 'Årsforbruk for artikkel i (enheter per år)', '1 000 stk/år'],
        ['UNIT_PRICE_i', 'Pris per enhet for artikkel i (kr/stk)', 'kr 50/stk'],
    ]
)
add_feynman_tip(
    'Hvis du bruker 1000 hansker i året og hver hanske koster 50 kr, da er '
    'årsverdien for hansker 1000 x 50 = 50 000 kr. Så enkelt er det.'
)

add_heading('Formel: Kumulativ verdiandel', 3)
add_formula_box(
    'Formel:',
    'C_i = (v_1 + v_2 + ... + v_i) / V_tot',
    'Kumulativ verdiandel = summen av verdier fra artikkel 1 til i, delt på total verdi.'
)
add_table_simple(
    ['Symbol', 'Betydning'],
    [
        ['C_i', 'Kumulativ verdiandel etter artikkel i (mellom 0 og 1)'],
        ['V_tot', 'Total årsverdi for alle artikler samlet'],
    ]
)
add_para('Klassifiseringsregler:')
add_table_simple(
    ['Klasse', 'Regel', 'Tolkning'],
    [
        ['A', 'C_i <= 0.80 (80 %)', 'De mest verdifulle artiklene som utgjør 80 % av totalverdien'],
        ['B', '0.80 < C_i <= 0.95', 'Neste 15 % av totalverdien'],
        ['C', 'C_i > 0.95', 'Resten – de billigste 5 % av verdien, men ofte flest artikler'],
    ]
)
add_feynman_tip(
    'Tenk på Netflix-topplisten: De 20 % mest sette filmene står for 80 % av '
    'all visningstid. Tilsvarende: De 26 % dyreste sykehusvarene står for '
    '80 % av all pengebruk. Det er Pareto-prinsippet – «de vitale få og de '
    'trivielle mange».'
)

# ─── 3.2 XYZ ───
add_heading('3.2 XYZ-klassifisering (variasjonskoeffisient)', 2)

add_formula_box(
    'Formel:',
    'CV_i = sigma_i / mu_i',
    'Variasjonskoeffisienten = standardavvik delt på gjennomsnitt av månedlig forbruk.'
)
add_table_simple(
    ['Symbol', 'Betydning', 'Eksempel'],
    [
        ['CV_i', 'Variasjonskoeffisient for artikkel i (dimensjonsløs)', '0.35'],
        ['sigma_i', 'Standardavvik av månedlig forbruk', '35 stk'],
        ['mu_i', 'Gjennomsnittlig månedlig forbruk', '100 stk'],
    ]
)
add_para('Klassifiseringsregler:')
add_table_simple(
    ['Klasse', 'CV-grense', 'Tolkning'],
    [
        ['X (stabilt)', 'CV < 0.5', 'Forbruket er jevnt og forutsigbart – lett å planlegge'],
        ['Y (moderat)', '0.5 <= CV < 1.0', 'Noe variasjon, men fortsatt håndterbart'],
        ['Z (uregelmessig)', 'CV >= 1.0', 'Svært uforutsigbart forbruk – vanskelig å planlegge'],
    ]
)
add_feynman_tip(
    'Tenk på strømregningen din. Hvis den er ca. 1000 kr hver måned (+-150 kr), '
    'er CV = 150/1000 = 0.15 – svært stabil (X-klasse). Men hvis noen måneder '
    'er 200 kr og andre er 3000 kr, er CV mye høyere – det er en Z-vare. '
    'Sentrallageret (HVFS) liker X-varer fordi de kan planlegge leveransene.'
)

add_heading('Viktig detalj: Hvorfor CV og ikke bare standardavvik?', 3)
add_para(
    'Standardavviket alene sier lite. En vare med forbruk 10 000/mnd og '
    'standardavvik 1 000 er faktisk veldig stabil (CV = 0.1). Men en vare '
    'med forbruk 5/mnd og standardavvik 5 er veldig ustabil (CV = 1.0). '
    'CV normaliserer variasjonen slik at vi kan sammenligne varer med '
    'helt ulikt volum.'
)

# ─── 3.3 EOQ ───
add_heading('3.3 EOQ – Wilson-modellen (optimal bestillingsmengde)', 2)

add_formula_box(
    'Formel:',
    'Q* = sqrt(2 x D x S / H)',
    'Optimal bestillingsmengde = kvadratroten av (2 x årsforbruk x ordrekostnad / holdekostnad per enhet).'
)
add_table_simple(
    ['Symbol', 'Betydning', 'Verdi i rapporten'],
    [
        ['Q*', 'Optimal bestillingsmengde (enheter per ordre)', 'Beregnes per artikkel'],
        ['D', 'Årsforbruk (enheter per år)', 'Fra MSEG-data'],
        ['S', 'Ordrekostnad per bestilling (kr)', '750 kr'],
        ['H', 'Holdekostnad per enhet per år (kr)', 'h x UNIT_PRICE = 20 % x pris'],
        ['h', 'Holdesats (prosent av enhetspris)', '20 %'],
    ]
)

add_feynman_tip(
    'EOQ handler om en avveining: Hvis du bestiller sjelden, bestiller du '
    'store mengder og har mye varer på lager (dyr lagring). Hvis du bestiller '
    'ofte, bestiller du lite hver gang, men betaler mange ordrekostnader. '
    'EOQ er det perfekte balansepunktet.\n\n'
    'Analogi: Skal du handle middag en gang i uken (billigere per tur, men '
    'maten rekker kanskje ikke) eller hver dag (ferskere, men dyrere i '
    'transport)? EOQ finner den optimale handlefrekvensen.'
)

add_heading('Talleksempel', 3)
add_para('En vare med årsforbruk D = 1000 stk, pris = 100 kr/stk:')
add_para('H = 20 % x 100 = 20 kr/stk/år', indent_cm=1)
add_para('Q* = sqrt(2 x 1000 x 750 / 20) = sqrt(75 000) = 274 stk per ordre', indent_cm=1)
add_para('Optimal frekvens f* = D / Q* = 1000 / 274 = 3.65 ordrer/år', indent_cm=1)

add_heading('Optimal ordrefrekvens', 3)
add_formula_box(
    'Formel:',
    'f* = D / Q* = sqrt(D x H / (2 x S))',
    'Optimal ordrefrekvens = årsforbruk delt på optimal bestillingsmengde.'
)
add_feynman_tip(
    'Hvis optimal bestillingsmengde er 274 stk og du bruker 1000 stk/år, '
    'bør du bestille 1000/274 = 3.65 ganger i året, altså ca. hver tredje måned.'
)

# ─── 3.4 Avviksformel ───
add_heading('3.4 Avviksformel (frekvensavvik)', 2)

add_formula_box(
    'Formel:',
    'FREQ_AVVIK_i = (f_obs,i - f*_i) / f*_i',
    'Relativt frekvensavvik = (faktisk ordrefrekvens minus optimal) delt på optimal.'
)
add_table_simple(
    ['Symbol', 'Betydning'],
    [
        ['FREQ_AVVIK_i', 'Relativt frekvensavvik for artikkel i'],
        ['f_obs,i', 'Faktisk observert ordrefrekvens (fra SAP-data)'],
        ['f*_i', 'Optimal ordrefrekvens (fra EOQ)'],
    ]
)

add_para('Terskelverdi:', bold=True)
add_para(
    'Terskelen er satt til tau_f = 1.5. Hvis FREQ_AVVIK > 1.5, betyr det at '
    'artikkelen bestilles mer enn 150 % oftere enn optimalt, og klassifiseres '
    'som FOR_MANGE_ORDRER.'
)

add_heading('Talleksempel', 3)
add_para('Optimal frekvens f* = 4 ordrer/år. Faktisk f_obs = 12 ordrer/år.', indent_cm=1)
add_para('FREQ_AVVIK = (12 - 4) / 4 = 2.0', indent_cm=1)
add_para('2.0 > 1.5 => FOR_MANGE_ORDRER!', indent_cm=1)

add_feynman_tip(
    'Artikkelen bestilles 3 ganger så ofte som den burde. Hver bestilling '
    'koster 750 kr i administrasjon, så det kastes bort 8 x 750 = 6000 kr/år '
    'i unødvendige ordrekostnader. Det er dette rapporten prøver å fikse.'
)

# ─── 3.5 Totalkostnad ───
add_heading('3.5 Totalkostnadsfunksjon', 2)

add_formula_box(
    'Formel:',
    'TC(f) = f x S + (D / (2f)) x H',
    'Total kostnad = ordrekostnader + holdekostnader.'
)
add_table_simple(
    ['Ledd', 'Formel', 'Hva det betyr'],
    [
        ['Ordrekostnader', 'f x S', 'Antall ordrer per år x kostnad per ordre'],
        ['Holdekostnader', '(D / 2f) x H', 'Gjennomsnittlig lagerbeholdning x holdekostnad'],
    ]
)

add_heading('Kostnadsavvik per artikkel', 3)
add_formula_box(
    'Formel:',
    'delta_TC_i = TC(f_obs,i) - TC(f*_i)',
    'Kostnadsavvik = totalkostnad ved faktisk frekvens minus totalkostnad ved optimal frekvens.'
)
add_feynman_tip(
    'Hvis du bestiller 12 ganger i året men burde bestille 4 ganger, '
    'betaler du for mye i ordrekostnader (8 ekstra x 750 kr) men sparer litt '
    'på lagerkostnader (lavere beholdning). delta_TC er nettodifferansen – '
    'det du taper totalt sett på å bestille feil.'
)

# ─── 3.6 Besparelse ───
add_heading('3.6 Besparelsesformel', 2)

add_formula_box(
    'Formel:',
    'B_HVFS = SUM(delta_TC_i) x g    (for alle artikler i OVERFØR)',
    'Total besparelse = sum av kostnadsavvik for alle HVFS-artikler, ganget med realiseringsgrad.'
)
add_table_simple(
    ['Symbol', 'Betydning', 'Verdier i rapporten'],
    [
        ['B_HVFS', 'Estimert årlig besparelse (kr)', 'Ca. 451 515 kr (base case)'],
        ['delta_TC_i', 'Kostnadsavvik per artikkel i (kr/år)', 'Varierer per artikkel'],
        ['g', 'Gevinstrealiseringsgrad (0–1)', '0.50 / 0.75 / 1.00'],
    ]
)
add_para('Tre scenarier:')
add_table_simple(
    ['Scenario', 'g', 'Besparelse'],
    [
        ['Worst case (pessimistisk)', '50 %', 'kr 301 010/år'],
        ['Base case (realistisk)', '75 %', 'kr 451 515/år'],
        ['Best case (optimistisk)', '100 %', 'kr 602 020/år'],
    ]
)
add_feynman_tip(
    'g = 75 % betyr at vi forventer å realisere tre fjerdedeler av den '
    'teoretiske besparelsen. De siste 25 % går tapt på grunn av '
    'implementeringsfriksjon – ting tar tid, leverandører kan endre '
    'vilkår, og SAP-systemet må omprogrammeres. Ingen gevinst er 100 % gratis.'
)

add_para(
    'Viktig: Kun 117 av de 145 OVERFØR-artiklene inngår i besparelses-'
    'beregningen – de 28 som anbefales overført basert på K-means alene '
    '(uten FOR_MANGE_ORDRER-flagg) har ikke et kvantifisert kostnadsavvik.',
    italic=True
)

# ─── 3.7 K-means feature ───
add_heading('3.7 K-means featurevektor', 2)

add_formula_box(
    'Formel:',
    'x_i = [ z(ln(CV_i)),  z(ln(v_i + 1)),  z(ln(|delta_TC_i| + 1)) ]',
    'Hver artikkel representeres som et punkt i 3D-rom med tre verdier.'
)
add_table_simple(
    ['Feature', 'Hva den måler', 'Transformasjoner'],
    [
        ['z(ln(CV_i))', 'Forbruksstabilitet', 'Logaritme, deretter z-score'],
        ['z(ln(v_i + 1))', 'Artikkelverdi (kr)', 'Logaritme (+1 for null), z-score'],
        ['z(ln(|delta_TC_i| + 1))', 'Kostnadsavvik (kr)', 'Absolutt verdi, log, z-score'],
    ]
)

add_heading('Hvorfor ln (logaritme)?', 3)
add_para(
    'Mange varer har lav verdi og noen få har svært høy verdi. Uten '
    'logaritme ville de dyre artiklene dominere hele analysen. Logaritmen '
    '«komprimerer» de store tallene slik at forskjellene mellom billige '
    'varer også synes.'
)
add_feynman_tip(
    'Tenk på Richter-skalaen for jordskjelv: Et skjelv på 7.0 er ikke '
    '«litt sterkere» enn 6.0 – det er 10 ganger sterkere. Skalaen er '
    'logaritmisk for å håndtere store variasjoner. Samme prinsipp her.'
)

add_heading('Hvorfor z-score?', 3)
add_para(
    'z-score standardiserer: z = (x - gjennomsnitt) / standardavvik. '
    'Etter z-score har alle tre features gjennomsnitt 0 og standardavvik 1. '
    'Uten dette ville den variabelen med de største tallene (f.eks. verdi i kr) '
    'dominere K-means-algoritmen.'
)

add_heading('Hvorfor +1 i ln(v + 1)?', 3)
add_para(
    'ln(0) er udefinert (minus uendelig). Noen artikler kan ha verdi = 0 '
    'eller delta_TC = 0. Ved å legge til 1 unngår vi ln(0). For store tall '
    'gjør +1 nesten ingen forskjell (ln(100001) ≈ ln(100000)).'
)

# ─── 3.8 Silhouette ───
add_heading('3.8 Silhouette-score', 2)

add_formula_box(
    'Formel:',
    's_i = (b_i - a_i) / max(a_i, b_i)',
    'Silhouette-score for punkt i = (avstand til naboklynge minus avstand innad) / den største.'
)
add_table_simple(
    ['Symbol', 'Betydning'],
    [
        ['s_i', 'Silhouette-score for punkt i (mellom -1 og +1)'],
        ['a_i', 'Gjennomsnittlig avstand til andre punkter i SAMME klynge'],
        ['b_i', 'Gjennomsnittlig avstand til punkter i NÆRMESTE andre klynge'],
    ]
)
add_para('Tolkning:')
add_table_simple(
    ['Score', 'Betydning'],
    [
        ['Nær +1', 'Punktet ligger tydelig i sin klynge, langt fra naboklyngen'],
        ['Nær 0', 'Punktet ligger på grensen mellom to klynger'],
        ['Negativ', 'Punktet er sannsynligvis i feil klynge'],
        ['> 0.3 (snitt)', 'Akseptabel klyngestruktur for eksplorativ analyse'],
    ]
)
add_feynman_tip(
    'Tenk på elever i en skolegård. Hvis alle guttene står i én gruppe og '
    'alle jentene i en annen, har gruppene høy silhouette-score – alle er '
    'nær «sine egne» og langt fra «de andre». Hvis alle blandes tilfeldig, '
    'er scoren nær null. Rapportens score på 0.383 betyr at klyngene er '
    'tydelige nok til å være nyttige, men ikke perfekt separerte.'
)

add_para(
    'I rapporten: Silhouette-score beregnes for K = 2, 3, 4, 5, 6, 7. '
    'K = 3 ga høyest score (0.383 på treningsdata, 0.368 på testdata). '
    'At testscoren er nær treningsscoren viser at klyngene ikke er '
    'et artefakt av treningsdataene (dvs. modellen generaliserer).',
    italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DEL 4 – ALLE FIGURER FORKLART
# ══════════════════════════════════════════════════════════════════
add_heading('Del 4 – Alle figurer forklart (Figur 1–12)', 1)

add_para(
    'Rapporten inneholder 12 figurer. Her forklares hva hver figur viser, '
    'hvordan den skal leses, og hva den betyr i vanlige ord.'
)

figures = [
    ('Figur 1 – Konseptuelt rammeverk (Fig00)',
     'plot_rammeverk.py -> Fig00_Konseptuelt_Rammeverk.png',
     'Hva den viser:',
     'Et flytdiagram over hele analyseprosessen – fra rådata i SAP til ferdig '
     'HVFS-anbefaling. Den viser de fire analysemetodene (ABC, XYZ, EOQ, K-means) '
     'som parallelle strømmer som samles i regelmotoren.',
     'Slik leser du den:',
     'Start øverst med SAP-data og følg pilene nedover. Legg merke til at '
     'ABC og XYZ er uavhengige (beregnes parallelt), mens regelmotoren til slutt '
     'kombinerer alt til én anbefaling per artikkel.'),

    ('Figur 2 – Lagerstruktur (Fig01)',
     'plot_lagerstruktur.py -> Fig01_Lagerstruktur.png',
     'Hva den viser:',
     'Et organisasjonskart over forsyningskjeden i Helse Vest: Helse Vest RHF '
     'øverst, helseforetakene (Bergen, Stavanger, etc.) i midten, og HVFS som '
     'regionalt sentrallager.',
     'Slik leser du den:',
     'I dag: Helse Bergen bestiller direkte fra leverandører til sitt lokale lager. '
     'Etter HVFS: Mange varer bestilles sentralt via HVFS og leveres '
     'avdelingspakket (APL).'),

    ('Figur 3 – Analysepipeline (Fig02)',
     'plot_analysepipeline.py -> Fig02_Analysepipeline.png',
     'Hva den viser:',
     'Den tekniske datapipelinen: fra SAP-tabeller (MARD, MSEG, EKKO...) '
     'gjennom dataforbehandling (D-01 til D-08) til MASTERFILE og videre '
     'til de fire analysene.',
     'Slik leser du den:',
     'Følg dataflyten fra venstre til høyre. Legg merke til at 1006 artikler '
     'kommer inn og 709 kommer ut etter filtrering (D-01).'),

    ('Figur 4 – Regelmotor (Fig03)',
     'plot_regelmotor.py -> Fig03_Regelmotor.png',
     'Hva den viser:',
     'Et flytskjema over de 8 beslutningsreglene (R1–R8). Hver rute er en '
     'betingelse (f.eks. «Er XYZ = Z?»), og pilene viser hva som skjer '
     'ved ja/nei.',
     'Slik leser du den:',
     'Start øverst. R1 spør om artikkelen er Z-klassifisert. Hvis ja -> '
     'BEHOLD LOKALT (uforutsigbart forbruk). Hvis nei, gå til R2, osv. '
     'Den første regelen som treffer, bestemmer utfallet.'),

    ('Figur 5 – Dataoversikt (Fig04)',
     'plot_dataoversikt.py -> Fig04_Dataoversikt.png',
     'Hva den viser:',
     'Histogrammer og fordelingsplott for nøkkelvariabler i datasettet: '
     'enhetspris, årsforbruk, CV, lagerbeholdning etc.',
     'Slik leser du den:',
     'Hvert subplot viser fordelingen av én variabel. Lange høyre haler '
     '(skjevhet) er typisk – noen få artikler har svært høye verdier.'),

    ('Figur 6 – ABC Pareto-kurve (Fig05)',
     'plot_abc_pareto.py -> Fig05_ABC_Pareto.png',
     'Hva den viser:',
     'En graf der x-aksen er artikler (sortert fra dyreste til billigste) '
     'og y-aksen er kumulativ verdiandel (0–100 %). To horisontale linjer '
     'ved 80 % og 95 % markerer A/B- og B/C-grensene.',
     'Slik leser du den:',
     'Kurven stiger bratt først (få A-artikler bærer mye verdi) og flater '
     'deretter ut (mange C-artikler bærer lite verdi). Skjæringspunktene '
     'med 80 %- og 95 %-linjene avgrenser A, B og C.'),

    ('Figur 7 – ABC/XYZ-matrise (Fig06)',
     'plot_abc_xyz_matrise.py -> Fig06_ABC_XYZ_Matrise.png',
     'Hva den viser:',
     'En 3x3-rutenett med ABC-klassene på én akse og XYZ på den andre. '
     'Hvert felt viser antall artikler i den kombinasjonen (f.eks. AX = 120).',
     'Slik leser du den:',
     'AX (øverst til venstre) = høy verdi, stabilt forbruk – de beste '
     'kandidatene for sentralisering. CZ (nederst til høyre) = lav verdi, '
     'uforutsigbart – beholdes lokalt. Fargekoding: grønt = bra for HVFS, '
     'rødt = behold lokalt.'),

    ('Figur 8 – EOQ-avvik (Fig07)',
     'plot_eoq_avvik.py -> Fig07_EOQ_Avvik.png',
     'Hva den viser:',
     'Et scatterplott eller histogram som viser FREQ_AVVIK for alle artikler. '
     'En vertikal linje ved tau_f = 1.5 skiller «OK» fra «FOR_MANGE_ORDRER».',
     'Slik leser du den:',
     'Artikler til høyre for den røde linjen bestiller for ofte. Jo lenger '
     'til høyre, jo mer suboptimal er bestillingspraksisen. 73 % av artiklene '
     'er til høyre for linjen.'),

    ('Figur 9 – Silhouette-score (Fig08)',
     'plot_silhouette.py -> Fig08_Silhouette.png',
     'Hva den viser:',
     'Et linjediagram med K (antall klynger) på x-aksen og gjennomsnittlig '
     'silhouette-score på y-aksen, for K = 2 til 7.',
     'Slik leser du den:',
     'Den høyeste punktet er ved K = 3 (score 0.383). Det betyr at 3 klynger '
     'gir den tydeligste gruppestrukturen i dataene.'),

    ('Figur 10 – K-means klynger (Fig09)',
     'plot_kmeans_klynger.py -> Fig09_Kmeans_Klynger.png',
     'Hva den viser:',
     'Et scatterplott der hver prikk er en artikkel, plottet langs to av '
     'de tre featurene (CV og verdi). Fargene viser hvilken klynge artikkelen '
     'tilhører. K_OVERFØR-klyngen er markert med en stjerne.',
     'Slik leser du den:',
     'Grønne prikker (K_OVERFØR) = artikler med lav CV og høy verdi – de '
     'beste HVFS-kandidatene. Separasjon mellom fargene = klyngene er tydelige.'),

    ('Figur 11 – Klyngeprofiler (Fig10)',
     'plot_kmeans_profil.py -> Fig10_Kmeans_Profil.png',
     'Hva den viser:',
     'Et linjediagram der x-aksen er de tre featurene og y-aksen er '
     'gjennomsnittlig z-score per klynge. Tre linjer – en per klynge.',
     'Slik leser du den:',
     'K_OVERFØR-klyngen (grønn) har lav CV-score (stabil) og høy verdi-score '
     '(dyr). De andre klyngene har motsatte profiler.'),

    ('Figur 12 – Regelmotor og besparelse (Fig11)',
     'plot_regelmotor_besparelse.py -> Fig11_Regelmotor_Besparelse.png',
     'Hva den viser:',
     'To paneler: (1) Søylediagram med fordeling av anbefalinger '
     '(OVERFØR/BEHOLD/VURDER/MANGLER). (2) Søylediagram med estimert '
     'besparelse for de tre scenariene (worst/base/best).',
     'Slik leser du den:',
     'Venstre panel: VURDER NÆRMERE er den største gruppen (284) – mange '
     'artikler trenger manuell gjennomgang. Høyre panel: Base case gir '
     'ca. 450 000 kr/år i besparelse.'),
]

for title, script, label1, text1, label2, text2 in figures:
    add_heading(title, 2)
    add_para(f'Kilde: {script}', italic=True, color=GREY)
    add_para(label1, bold=True)
    add_para(text1, indent_cm=0.5)
    add_para(label2, bold=True)
    add_para(text2, indent_cm=0.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DEL 5 – ALLE TABELLER FORKLART
# ══════════════════════════════════════════════════════════════════
add_heading('Del 5 – Alle tabeller forklart (Tabell 1–15)', 1)

tables = [
    ('Tabell 1 – Litteraturoversikt',
     'Viser de 23 akademiske kildene rapporten bygger på, organisert etter '
     'hovedtema. For eksempel: Volland et al. (2017) handler om '
     'materiallogistikk i sykehus og er relevant fordi de viser at lagerstyring '
     'utgjør 30–40 % av driftskostnadene.',
     'Denne tabellen er rapportens «kildekart» – den viser hvor '
     'hvert faglig argument kommer fra.'),

    ('Tabell 2 – Sammenligning av analysemetoder',
     'Sammenligner ABC, XYZ, EOQ og K-means langs fire dimensjoner: styrker, '
     'svakheter, forutsetninger og referanse.',
     'En «hurtigoppslagstabell» for å forstå trade-off-ene mellom metodene.'),

    ('Tabell 3 – Nøkkeltall for casevirksomheten',
     'Oppsummerer hva Helse Bergen er i SAP-termer: WERKS 3300, LGORT 3001, '
     '709 aktive artikler, 14 SAP-tabeller, innkjøpsgruppe 300/3000.',
     'Alle tall som identifiserer caset – tilsvarende «pasientjournalen» for lageret.'),

    ('Tabell 4 – Datagrunnlag: 14 SAP-tabeller',
     'Lister alle 14 SAP-tabeller med beskrivelse og funksjonell kategori '
     '(masterdata, forbruksdata, innkjøpsdata, supplerende). For eksempel: '
     'MSEG inneholder varebevegelser (forbruk til avdeling), EKPO inneholder '
     'innkjøpsordreposisjoner.',
     'Dataspesifikasjonen – viser eksakt hvor tallene kommer fra i SAP.'),

    ('Tabell 5 – Datavalgsbeslutninger D-01–D-08',
     'Dokumenterer åtte valg som ble tatt under databehandlingen. For eksempel: '
     'D-01 (ekskluder inaktive artikler: 1006 -> 709), D-02 (korriger enhetspris '
     'for prisenhet: STPRS / PEINH), D-04 (erstatt SAP ZZXYZ med beregnet CV).',
     'Rapporten er transparent: Enhver analytiker kan se nøyaktig hvilke valg '
     'som er tatt og hvorfor. Ingen «svarte bokser».'),

    ('Tabell 6 – Modellparametere',
     'Alle parameterinnstillinger samlet: ABC-grenser (80/95 %), XYZ-grenser '
     '(CV 0.5/1.0), S = 750 kr, h = 20 %, K-means-innstillinger, random_state = 42.',
     'Parameterkortet – alle «bryterne» i analysen på ett sted.'),

    ('Tabell 7 – Regelmotor: 8 beslutningsregler',
     'De 8 reglene R1–R8 med betingelse, anbefaling og logikk. R1: Z-artikler '
     '-> BEHOLD. R3: A/B + X + FOR_MANGE -> OVERFØR. Osv.',
     'Selve «oppskriften» for regelmotoren. Reglene kjøres i rekkefølge – '
     'den første som treffer, avgjør.'),

    ('Tabell 8 – ABC-fordeling',
     '182 A-artikler (25.7 %, 80 % av verdien), 184 B (26.0 %), '
     '338 C (47.7 %), 5 ikke klassifisert.',
     'Bekrefter Pareto-mønsteret: ca. en fjerdedel av artiklene bærer 80 % av verdien.'),

    ('Tabell 9 – XYZ-fordeling',
     '350 X (50.9 %, stabilt), 193 Y (28.1 %, moderat), 144 Z (20.9 %, '
     'uregelmessig), 22 ikke klassifisert.',
     'Over halvparten av artiklene har stabilt forbruk – et godt tegn for HVFS.'),

    ('Tabell 10 – SAP ZZXYZ-validering',
     'Krysstabell mellom SAP-systemets ZZXYZ-klasse og den beregnede CV-klassen. '
     'Kun 33 % samsvar (125 av 375 artikler).',
     'Viser at SAP-systemets XYZ-felt er utdatert og upålitelig – '
     'beregnet CV bør erstatte det.'),

    ('Tabell 11 – EOQ-avviksresultater',
     '356 FOR_MANGE_ORDRER (73.1 %), 100 OK (20.5 %), 31 FOR_FÅ_ORDRER (6.4 %). '
     'Samlet teoretisk ΔTC = kr 2 333 441/år for alle 487 analyserte artikler.',
     'Tre av fire artikler bestilles oftere enn EOQ tilsier. '
     'Merk: kr 2,3 mill er det teoretiske samlede kostnadsavviket for HELE '
     'populasjonen – det er IKKE det samme som realiserbar besparelse. '
     'Faktisk besparelsesestimat (Tabell 14) gjelder kun de 117 artiklene '
     'som anbefales overført OG har FOR_MANGE_ORDRER: kr 301 010–602 020/år, '
     'base case kr 451 515.'),

    ('Tabell 12 – K-means klyngeprofiler',
     'Tre klynger med gjennomsnittsverdier: Klynge 1 (31 stk, lav verdi, '
     'høy CV), Klynge 2 (175 stk, middels verdi, svært variabelt), '
     'Klynge 3/K_OVERFØR (281 stk, høy verdi, stabil, høyt avvik).',
     'K_OVERFØR-klyngen er de naturlige HVFS-kandidatene: verdifulle, '
     'stabile og overbestilte.'),

    ('Tabell 13 – HVFS-anbefalinger',
     '145 OVERFØR (20.5 %), 257 BEHOLD (36.2 %), 284 VURDER (40.1 %), '
     '23 MANGLER DATA (3.2 %).',
     'Hovedresultatet i rapporten – den endelige «listen» over anbefalinger.'),

    ('Tabell 14 – Besparelsesestimater',
     'Worst: kr 301 010 (g=50 %), Base: kr 451 515 (g=75 %), '
     'Best: kr 602 020 (g=100 %). For 117 artikler med S = 750 kr.',
     'Svaret på «hvor mye kan vi spare?» – mellom 300k og 600k kr/år.'),

    ('Tabell 15 – Egne resultater vs. litteratur',
     'Sammenstilling av 6 nøkkelfunn mot referanser: A-klasse 25.7 % '
     '(OK vs. Pareto), ZZXYZ 33 % (forventet), silhouette 0.383 (> 0.3), osv.',
     'Alt stemmer med litteraturen – rapporten finner det man ville forvente.'),
]

for title, content, interpretation in tables:
    add_heading(title, 2)
    add_para('Innhold:', bold=True)
    add_para(content, indent_cm=0.5)
    add_para('Tolkning:', bold=True)
    add_para(interpretation, indent_cm=0.5, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DEL 6 – ORDLISTE
# ══════════════════════════════════════════════════════════════════
add_heading('Del 6 – Ordliste og nøkkelbegreper', 1)

glossary = [
    ('ABC-analyse', 'Sortering av lagervarer etter verdi. A = de dyreste (80 % av total verdi), '
     'B = neste 15 %, C = resten (billigst, men flest varer).'),
    ('APL', 'Avdelingspakkede leveranser – ferdigpakkede kasser direkte til sykehusavdelingen.'),
    ('BWART', 'Bevegelsestype i SAP. 201 = forbruk til kostnadssted. 647 = spesialforbruk. '
     '101 = varemottak fra leverandør.'),
    ('CV', 'Variasjonskoeffisient = standardavvik / gjennomsnitt. Måler relativ variasjon.'),
    ('D_ANNUAL', 'Annualisert årsforbruk (enheter per år).'),
    ('delta_TC', 'Kostnadsavvik mellom faktisk og optimal totalkostnad (kr/år).'),
    ('EKBE', 'SAP-tabell for innkjøpsordrehistorikk (varemottak).'),
    ('EKKO/EKPO', 'SAP-tabeller for innkjøpsordrehode og -posisjoner.'),
    ('EOQ', 'Economic Order Quantity – optimal bestillingsmengde som minimerer totalkostnad.'),
    ('f*', 'Optimal ordrefrekvens (ordrer per år) ifølge Wilson-modellen.'),
    ('f_obs', 'Observert (faktisk) ordrefrekvens fra SAP-data.'),
    ('FREQ_AVVIK', 'Relativt avvik mellom faktisk og optimal ordrefrekvens.'),
    ('g', 'Gevinstrealiseringsgrad – andelen av teoretisk besparelse som faktisk oppnås.'),
    ('H', 'Holdekostnad per enhet per år = h x UNIT_PRICE (kr/stk/år).'),
    ('h', 'Holdesats – prosent av enhetspris som utgjør årlig holdekostnad. h = 20 %.'),
    ('HVFS', 'Helse Vest Forsyningssenter – det regionale sentrallageret.'),
    ('K-means', 'Klyngealgoritme som deler datapunkter i K grupper basert på likhet.'),
    ('K_OVERFØR', 'Den K-means-klyngen som best matcher HVFS-profilen (lav CV, høy verdi).'),
    ('LGORT', 'Lagersted i SAP. LGORT 3001 = Helse Bergens hovedlager.'),
    ('LIBRA', 'Helse Vests regionale SAP S/4HANA-implementeringsprosjekt.'),
    ('ln', 'Naturlig logaritme. Komprimerer store tall for bedre statistisk analyse.'),
    ('MASTERFILE', 'Den konsoliderte analysefilen med 709 artikler og 21 kolonner.'),
    ('MSEG', 'SAP-tabell for varebevegelser (forbruk, overføring, retur).'),
    ('PEINH', 'Prisenhet i SAP. Hvis PEINH = 100, er oppgitt pris per 100 stk, ikke per stk.'),
    ('Q*', 'Optimal bestillingsmengde ifølge Wilson/EOQ-formelen.'),
    ('S', 'Ordrekostnad per bestilling. S = 750 NOK i denne studien.'),
    ('SE16H', 'SAP-transaksjon for direkte lesetilgang til databasetabeller.'),
    ('Silhouette-score', 'Mål på klyngekvalitet. +1 = perfekt separert. 0 = overlapp. '
     '-1 = feil klynge.'),
    ('STPRS', 'Standardpris i SAP (fra tabell MBEW). Oppgitt per PEINH enheter.'),
    ('tau_f', 'Terskelverdi for EOQ-frekvensavvik. tau_f = 1.5 -> 150 % overbestilling.'),
    ('UNIT_PRICE', 'Korrigert enhetspris = STPRS / PEINH (kr per stk).'),
    ('WERKS', 'Anlegg/verk i SAP. WERKS 3300 = Helse Bergen.'),
    ('Wilson-modellen', 'Den klassiske EOQ-formelen: Q* = sqrt(2DS/H).'),
    ('z-score', 'Standardisert verdi: z = (x - gjennomsnitt) / standardavvik. '
     'Gjør ulike variabler sammenlignbare.'),
    ('ZZXYZ', 'SAP-felt for XYZ-klassifisering i tabell MDMA. Ofte utdatert.'),
]

add_table_simple(
    ['Begrep', 'Forklaring'],
    glossary
)

# ─── SISTE SIDE ───
doc.add_page_break()
add_heading('Om dette dokumentet', 1)
add_para(
    'Dette hjelpedokumentet er laget med Feynman-teknikken: Alle konsepter, '
    'formler, figurer og tabeller i LOG650-rapporten er forklart med enkle '
    'ord, analogier og talleksempler – som om du forklarer det til noen '
    'uten forkunnskaper i statistikk eller lagerstyring.'
)
add_para(
    'Feynman-teknikken har fire steg:\n'
    '1. Velg et konsept\n'
    '2. Forklar det med enkle ord (som til en 12-åring)\n'
    '3. Identifiser hullene i forklaringen din\n'
    '4. Forenkle og bruk analogier'
)
add_para(
    'Referanse: Basert på rapporten «Fra lokalt forsyningslager til '
    'regional sentralforsyning» av Thomas Ekrem Jensen, LOG650 '
    'Forskningsprosjekt, Høgskolen i Molde, våren 2026.',
    italic=True
)

# ─── LAGRE ───
doc.save(OUTPUT)
print(f'\nFerdig! Dokumentet er lagret som: {OUTPUT}')
print(f'Antall sider (ca.): {len(doc.paragraphs) // 25}+')
