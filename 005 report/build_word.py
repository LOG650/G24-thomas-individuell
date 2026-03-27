"""
Bygg Word-dokument frå HiMolde-malen + MD-rapport.
Opnar malen, fyller inn forsider, slettar plasshaldarar,
og legg til alt innhald frå MD-fila.
"""

import re, os, sys
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

import latex2mathml.converter
from lxml import etree

TEMPLATE = '../000 templates/Mal prosjekt LOG650 v2.docx'
_XSLT_CANDIDATES = [
    r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL',
    r'C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL',
    r'C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL',
    r'C:\Program Files\Microsoft Office\Office15\MML2OMML.XSL',
]

XSLT_PATH = None
for _candidate in _XSLT_CANDIDATES:
    if os.path.isfile(_candidate):
        XSLT_PATH = _candidate
        break
if XSLT_PATH is None:
    sys.exit(
        'FEIL: Fann ikkje MML2OMML.XSL. '
        'Legg til rett sti i _XSLT_CANDIDATES i build_word.py.\n'
        f'Søkte i: {_XSLT_CANDIDATES}'
    )

# Førebu XSLT-transformasjon for matte
_xslt = etree.parse(XSLT_PATH)
_math_transform = etree.XSLT(_xslt)
MD_FILE = 'LOG650_Rapport_FINAL.md'
OUTPUT = 'LOG650_Rapport.docx'

doc = Document(TEMPLATE)

# ════════════════════════════════════════════════════
# STEG 1: Fyll inn forsidefelter
# ════════════════════════════════════════════════════
print("Steg 1: Fyller inn forsider...")

field_map = {
    'Tittel (norsk og/eller engelsk)': 'Fra lokalt forsyningslager til regional sentralforsyning:\nMultikriterieklassifisering og klyngeanalyse for identifisering av overføringskandidater ved Helse Bergen',
    'Forfatter(e)': 'Thomas Ekrem Jensen',
    'Molde, Innleveringsdato': 'Stavanger, Mai 2026',
}

for para in doc.paragraphs:
    t = para.text.strip()
    for key, val in field_map.items():
        if t == key or t.startswith(key):
            para.clear()
            run = para.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            break
    # Studiepoeng og veileder (delvis match)
    if t.startswith('Studiepoeng:'):
        para.clear()
        r = para.add_run('Studiepoeng: 15')
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    elif t.startswith('Veileder:'):
        para.clear()
        r = para.add_run('Veileder: Bård Inge Austigard Pettersen')
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    elif t.startswith('Dato:') and len(t) < 20:
        para.clear()
        r = para.add_run('Dato: Mai 2026')
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    elif t.startswith('Antall ord:'):
        para.clear()
    elif t.startswith('Forfattererklæring:'):
        para.clear()

# ════════════════════════════════════════════════════
# STEG 2: Les MD-fila
# ════════════════════════════════════════════════════
print("Steg 2: Les MD-fil...")

with open(MD_FILE, 'r', encoding='utf-8') as f:
    md_all = f.read()

# Hent ut Forord-tekst
forord_m = re.search(r'^# Forord\s*\n(.*?)(?=\n---)', md_all, re.DOTALL | re.MULTILINE)
forord_text = forord_m.group(1).strip() if forord_m else ''

# Hent ut Sammendrag-tekst
sammendrag_m = re.search(r'^# Sammendrag\s*\n(.*?)(?=\n---)', md_all, re.DOTALL | re.MULTILINE)
sammendrag_text = sammendrag_m.group(1).strip() if sammendrag_m else ''

# ════════════════════════════════════════════════════
# STEG 3: Sett inn Sammendrag-tekst etter "Sammendrag"-overskrifta
# ════════════════════════════════════════════════════
print("Steg 3: Setter inn Forord og Sammendrag...")

def insert_text_after(anchor_para, text):
    """Set inn avsnitt etter eit gitt avsnitt."""
    current = anchor_para._element
    for chunk in text.split('\n\n'):
        chunk = chunk.strip()
        if not chunk:
            continue
        new_p = OxmlElement('w:p')
        # Stil
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Normal')
        pPr.append(pStyle)
        # Linjeavstand 1.5
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        new_p.append(pPr)
        # Tekst
        new_r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')  # 12pt = 24 half-points
        rPr.append(sz)
        new_r.append(rPr)
        new_t = OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        # Fjern markdown-formatering for enkel innsetting
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', chunk)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = re.sub(r'\$(.+?)\$', r'\1', clean)
        new_t.text = clean
        new_r.append(new_t)
        new_p.append(new_r)
        current.addnext(new_p)
        current = new_p

# Finn Sammendrag og Abstract i malen
for para in doc.paragraphs:
    if para.text.strip() == 'Sammendrag':
        insert_text_after(para, sammendrag_text)
    elif para.text.strip() == 'Abstract':
        # Sett inn ei enkel melding
        insert_text_after(para, '(Sammendrag på norsk – sjå avsnittet over.)')

# Sett inn Forord mellom "Antall ord"/"Forfattererklæring" og "Sammendrag"
# Finn sammendrag-paragrafen og set inn Forord-heading + tekst FØR den
for para in doc.paragraphs:
    if para.text.strip() == 'Sammendrag':
        # Lag Forord-heading
        forord_heading = OxmlElement('w:p')
        fh_pPr = OxmlElement('w:pPr')
        fh_pStyle = OxmlElement('w:pStyle')
        fh_pStyle.set(qn('w:val'), 'Heading1')
        fh_pPr.append(fh_pStyle)
        # Sideskift før
        fh_pb = OxmlElement('w:pageBreakBefore')
        fh_pb.set(qn('w:val'), 'true')
        fh_pPr.append(fh_pb)
        forord_heading.append(fh_pPr)
        fh_r = OxmlElement('w:r')
        fh_t = OxmlElement('w:t')
        fh_t.text = 'Forord'
        fh_r.append(fh_t)
        forord_heading.append(fh_r)
        para._element.addprevious(forord_heading)
        # Set inn forord-tekst etter heading
        insert_text_after_elem = forord_heading
        current = forord_heading
        for chunk in forord_text.split('\n\n'):
            chunk = chunk.strip()
            if not chunk:
                continue
            new_p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:line'), '360')
            sp.set(qn('w:lineRule'), 'auto')
            pPr.append(sp)
            new_p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), 'Times New Roman')
            rf.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rf)
            szz = OxmlElement('w:sz')
            szz.set(qn('w:val'), '24')
            rPr.append(szz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', chunk)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            t.text = clean
            r.append(t)
            new_p.append(r)
            current.addnext(new_p)
            current = new_p
        break

# ════════════════════════════════════════════════════
# STEG 4: Slett plasshaldar-kapittel (frå "Innledning" til slutt)
# ════════════════════════════════════════════════════
print("Steg 4: Slettar plasshaldarar...")

body = doc.element.body
to_remove = []
found = False
for child in list(body):
    if child.tag.endswith('}sectPr'):
        continue
    if not found:
        if child.tag.endswith('}p'):
            pPr_elem = child.find(qn('w:pPr'))
            ps = pPr_elem.find(qn('w:pStyle')) if pPr_elem is not None else None
            if ps is not None:
                val = ps.get(qn('w:val'), '')
                if 'Heading' in val or 'Overskrift' in val or val == 'Overskrift1':
                    txt = ''.join(t.text or '' for t in child.iter(qn('w:t')))
                    if 'Innledning' in txt:
                        found = True
                        to_remove.append(child)
    else:
        to_remove.append(child)

for elem in to_remove:
    body.remove(elem)

print(f"  Fjerna {len(to_remove)} plasshaldar-element")

# Fjern malens statiske TOC-oppføringar (toc 1, toc 2 osv.)
toc_remove = []
for para in list(doc.paragraphs):
    if para.style and para.style.name in ('toc 1', 'toc 2', 'toc 3', 'TOC Heading'):
        toc_remove.append(para._element)
for elem in toc_remove:
    body.remove(elem)
print(f"  Fjerna {len(toc_remove)} statiske TOC-oppføringar frå malen")

# ════════════════════════════════════════════════════
# STEG 5: Legg til kapittelinnhald
# ════════════════════════════════════════════════════
print("Steg 5: Legg til kapittelinnhald...")

# (Figurliste og Tabelliste er no auto-genererte Word-felt, sjå nedanfor)

# Finn start av kapittelinnhald i MD
kap_start = md_all.find('# Kapittel 1')
chapter_md = md_all[kap_start:]
lines = chapter_md.split('\n')

def latex_to_omml(latex_str):
    """Konverter LaTeX til OMML XML-element for Word."""
    try:
        # Rens LaTeX litt
        clean = latex_str.strip()
        if not clean:
            return None
        mathml = latex2mathml.converter.convert(clean)
        mathml_tree = etree.fromstring(mathml.encode('utf-8'))
        omml_tree = _math_transform(mathml_tree)
        return omml_tree.getroot()
    except Exception as e:
        print(f"  Math-feil: {e} for '{latex_str[:40]}'")
        return None

def add_math_paragraph(doc, latex_str):
    """Legg til ein sentrert matteblokk (display math)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.5
    omml = latex_to_omml(latex_str)
    if omml is not None:
        para._element.append(omml)
    else:
        # Fallback: vanleg tekst
        r = para.add_run(latex_str)
        r.font.name = 'Cambria Math'
        r.font.size = Pt(11)
    return para

def wrap_plain_math(text):
    """Pre-prosessor: wrap plain-text matematiske uttrykk i $...$."""
    # S = 750, h = 20, K = 3, g = 75, n = 487 osv.
    text = re.sub(r'(?<!\$)\b(S\s*=\s*\d[\d\s]*)', r'$\1$', text)
    text = re.sub(r'(?<!\$)\b(h\s*=\s*\d[\d\s,\.%]*%)', r'$\1$', text)
    text = re.sub(r'(?<!\$)\b(K\s*=\s*\d+)', r'$\1$', text)
    text = re.sub(r'(?<!\$)\b(g\s*=\s*\d[\d\s]*\s*%)', r'$\1$', text)
    text = re.sub(r'(?<!\$)\b(n\s*=\s*\d[\d\s]*)', r'$\1$', text)
    # CV < 0,5 etc.
    text = re.sub(r'(?<!\$)(CV\s*[<>≤≥]\s*[\d,]+)', r'$\1$', text)
    # τ_f = 1,5
    text = re.sub(r'(?<!\$)(τ_f\s*=\s*[\d,]+)', r'$\1$', text)
    # D_ANNUAL, UNIT_PRICE etc. som variabelnamn
    text = re.sub(r'(?<!\$)\b(D_ANNUAL\s*[=><]\s*\d+)', r'$\1$', text)
    # Fjern doble dollar-teikn som kan oppstå
    text = re.sub(r'\$\$([^$]+)\$\$', r'$$\1$$', text)
    return text

def set_run_font(run, name='Times New Roman', size=12):
    """Sett font på ein run."""
    run.font.name = name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    rf.set(qn('w:ascii'), name)
    rf.set(qn('w:hAnsi'), name)
    rf.set(qn('w:cs'), name)

def add_formatted_para(doc, text):
    """Legg til avsnitt med inline-formatering og inline-matte."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Erstatt em dash (—) med en dash (–)
    text = text.replace('—', '–')

    # Pre-prosesser: slå saman WORD$_subscript...$ til $\text{WORD}_subscript...$
    text = re.sub(
        r'([A-Za-z][A-Za-z0-9]*(?:\\?_[A-Za-z0-9]+)*)\$(_[^$]*)\$',
        lambda m: r'$\text{' + m.group(1).replace('\\', '').replace('_', r'\_') + '}' + m.group(2) + '$',
        text
    )

    # Pre-prosesser: wrap plain-text math i $...$
    text = wrap_plain_math(text)

    # Del opp i: inline code `...` | inline math $...$ | bold **...** | italic *...*
    segments = re.split(r'(`[^`]+?`|\$[^$]+?\$)', text)

    for seg in segments:
        if seg.startswith('`') and seg.endswith('`') and len(seg) > 2:
            # Inline code — Consolas monospace med grå bakgrunn
            code_text = seg[1:-1]
            r = para.add_run(code_text)
            r.font.size = Pt(10)
            # Sett font via rFonts-element (unngå duplikat)
            rpr = r._element.get_or_add_rPr()
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), 'Consolas')
            rf.set(qn('w:hAnsi'), 'Consolas')
            rf.set(qn('w:cs'), 'Consolas')
            rpr.insert(0, rf)
            # Grå bakgrunn
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F4F4F4')
            rpr.append(shd)
        elif seg.startswith('$') and seg.endswith('$') and len(seg) > 2:
            # Inline math — konverter til OMML
            latex = seg[1:-1]
            omml = latex_to_omml(latex)
            if omml is not None:
                para._element.append(omml)
                math_counter[0] += 1
            else:
                r = para.add_run(latex)
                r.font.name = 'Cambria Math'
                r.font.size = Pt(12)
                r.font.italic = True
        else:
            # Vanleg tekst — handter bold/italic
            parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', seg)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    r = para.add_run(part[2:-2])
                    r.font.bold = True
                    set_run_font(r)
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    r = para.add_run(part[1:-1])
                    r.font.italic = True
                    set_run_font(r)
                else:
                    r = para.add_run(part)
                    set_run_font(r)
    return para

def add_cell_content(para, text, font_size=11, bold=False):
    """Legg til celleinnhald med inline-matte ($...$) og bold (**...**)."""
    text = text.replace('—', '–')
    # Pre-prosesser: slå saman WORD$_subscript...$ til $\text{WORD}_subscript...$
    text = re.sub(
        r'([A-Za-z][A-Za-z0-9]*(?:\\?_[A-Za-z0-9]+)*)\$(_[^$]*)\$',
        lambda m: r'$\text{' + m.group(1).replace('\\', '').replace('_', r'\_') + '}' + m.group(2) + '$',
        text
    )
    # Del opp i inline math $...$ og resten
    segments = re.split(r'(\$[^$]+?\$)', text)
    for seg in segments:
        if seg.startswith('$') and seg.endswith('$') and len(seg) > 2:
            latex = seg[1:-1]
            omml = latex_to_omml(latex)
            if omml is not None:
                para._element.append(omml)
                math_counter[0] += 1
            else:
                r = para.add_run(latex)
                r.font.name = 'Cambria Math'
                r.font.size = Pt(font_size)
                r.font.italic = True
        else:
            # Rens escaped underscores og bold
            clean = seg.replace('\\', '').replace('—', '–')
            parts = re.split(r'(\*\*.*?\*\*)', clean)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    r = para.add_run(part[2:-2])
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(font_size)
                    r.font.bold = True
                else:
                    r = para.add_run(part)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(font_size)
                    r.font.bold = bold

def add_heading_formatted(doc, text, level):
    """Legg til overskrift med Times New Roman."""
    text = text.replace('—', '–')
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.insert(0, rf)
        rf.set(qn('w:ascii'), 'Times New Roman')
        rf.set(qn('w:hAnsi'), 'Times New Roman')
        if level == 1:
            r.font.size = Pt(18)
        elif level == 2:
            r.font.size = Pt(16)
        elif level == 3:
            r.font.size = Pt(14)
    return h

def set_threeline_borders(table):
    """Tre-linje tabellformat (booktabs):
    - Linje over header
    - Linje under header
    - Linje under siste rad
    - Ingen vertikale linjer, ingen bakgrunn
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Fjern eksisterande borders
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)

    # Tabellnivå: top og bottom = single, alt anna = nil
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:left w:val="nil"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:right w:val="nil"/>'
        '<w:insideH w:val="nil"/>'
        '<w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header-rad: legg til botn-linje på kvar celle
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            # Fjern eksisterande tcBorders
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            cell_borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                '<w:left w:val="nil"/>'
                '<w:right w:val="nil"/>'
                '</w:tcBorders>'
            )
            tcPr.append(cell_borders)

    # Fjern bakgrunnsfarge frå alle celler
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            for shd in tcPr.findall(qn('w:shd')):
                tcPr.remove(shd)

def add_field(paragraph, instr_text):
    """Set inn eit Word-felt (TOC, SEQ etc.) med fldChar begin/instrText/separate/end."""
    # Begin
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run_begin._element.append(fld_begin)
    # InstrText
    run_instr = paragraph.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    run_instr._element.append(instr)
    # Separate
    run_sep = paragraph.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run_sep._element.append(fld_sep)
    # Plasshaldartekst
    run_ph = paragraph.add_run('Oppdater felt: Ctrl+A, F9')
    run_ph.font.name = 'Times New Roman'
    run_ph.font.size = Pt(10)
    run_ph.font.color.rgb = RGBColor(128, 128, 128)
    # End
    run_end = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fld_end)

def add_seq_field_to_para(paragraph, seq_name, font_name='Times New Roman', font_size=10):
    """Set inn eit SEQ-felt inline i ein paragraf (for auto-nummerering av figurar/tabellar)."""
    # Begin
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run_begin._element.append(fld_begin)
    # InstrText
    run_instr = paragraph.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' SEQ {seq_name} \\* ARABIC '
    run_instr._element.append(instr)
    # Separate
    run_sep = paragraph.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run_sep._element.append(fld_sep)
    # Plasshaldarnummer
    run_num = paragraph.add_run('N')
    run_num.font.name = font_name
    run_num.font.size = Pt(font_size)
    run_num.font.italic = True
    # End
    run_end = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fld_end)

# Figurstiar
fig_map = {
    'Fig00': 'Fig00_Konseptuelt_Rammeverk.png',
    'Fig01': 'Fig01_Lagerstruktur.png',
    'Fig02': 'Fig02_Analysepipeline.png',
    'Fig03': 'Fig03_Regelmotor.png',
    'Fig04': 'Fig04_Dataoversikt.png',
    'Fig05': 'Fig05_ABC_Pareto.png',
    'Fig06': 'Fig06_ABC_XYZ_Matrise.png',
    'Fig07': 'Fig07_EOQ_Avvik.png',
    'Fig08': 'Fig08_Silhouette.png',
    'Fig09': 'Fig09_Kmeans_Klynger.png',
    'Fig10': 'Fig10_Kmeans_Profil.png',
    'Fig11': 'Fig11_Regelmotor_Besparelse.png',
}
plots_dir = os.path.abspath('../006 Analyse/plots')

# ── Innholdsfortegnelse (auto-generert Word TOC-felt) ──
print("  Legg til Innholdsfortegnelse...")
h = add_heading_formatted(doc, 'Innholdsfortegnelse', 1)
pPr = h._element.get_or_add_pPr()
pb = OxmlElement('w:pageBreakBefore')
pb.set(qn('w:val'), 'true')
pPr.append(pb)
toc_para = doc.add_paragraph()
toc_para.paragraph_format.line_spacing = 1.5
add_field(toc_para, r' TOC \o "1-3" \h \z \u ')

# ── Figurliste (auto-generert frå SEQ Figur-felt) ──
print("  Legg til Figurliste...")
h = add_heading_formatted(doc, 'Figurliste', 1)
pPr = h._element.get_or_add_pPr()
pb = OxmlElement('w:pageBreakBefore')
pb.set(qn('w:val'), 'true')
pPr.append(pb)
fig_toc_para = doc.add_paragraph()
fig_toc_para.paragraph_format.line_spacing = 1.5
add_field(fig_toc_para, r' TOC \h \z \c "Figur" ')

# ── Tabelliste (auto-generert frå SEQ Tabell-felt) ──
print("  Legg til Tabelliste...")
h = add_heading_formatted(doc, 'Tabelliste', 1)
pPr = h._element.get_or_add_pPr()
pb = OxmlElement('w:pageBreakBefore')
pb.set(qn('w:val'), 'true')
pPr.append(pb)
tab_toc_para = doc.add_paragraph()
tab_toc_para.paragraph_format.line_spacing = 1.5
add_field(tab_toc_para, r' TOC \h \z \c "Tabell" ')

print("  Legg til kapittel 1–9...")
table_count = 0
fig_count = 0
math_counter = [0]  # Liste for mutabilitet i nested scope
pending_caption = None  # Ubrukt — tabelltittel vert no sett inn FØR tabellen
in_referanseliste = False  # Flagg for APA 7 hengende innrykk
i = 0

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # ── Heading 1 ──
    if line.startswith('# ') and not line.startswith('## '):
        text = line[2:].strip()
        # Oppdater referanseliste-flagg for APA 7 hengende innrykk
        if text == 'Referanseliste':
            in_referanseliste = True
        elif in_referanseliste:
            in_referanseliste = False
        h = add_heading_formatted(doc, text, 1)
        # Sideskift før kapittel
        if any(kw in text for kw in ['Kapittel', 'Referanseliste', 'Vedlegg']):
            pPr = h._element.get_or_add_pPr()
            pb = OxmlElement('w:pageBreakBefore')
            pb.set(qn('w:val'), 'true')
            pPr.append(pb)

    # ── Heading 2 ──
    elif line.startswith('## '):
        add_heading_formatted(doc, line[3:].strip(), 2)

    # ── Heading 3 ──
    elif line.startswith('### '):
        add_heading_formatted(doc, line[4:].strip(), 3)

    # ── Horisontal linje ──
    elif stripped == '---':
        pass

    # ── Figur ──
    elif stripped.startswith('!['):
        m = re.match(r'!\[(.+?)\]\((.+?)\)', stripped)
        if m:
            caption = m.group(1).replace('—', '–')
            rel_path = m.group(2)
            # Finn biletet
            abs_path = os.path.abspath(os.path.join('.', rel_path))
            if os.path.exists(abs_path):
                doc.add_picture(abs_path, width=Cm(15))
                fig_count += 1
                # Bildetekst under med SEQ-felt for auto-nummerering
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_m = re.match(r'(Figur)\s+(\d+)\.\s*(.*)', caption)
                if cap_m:
                    r_label = cap_p.add_run('Figur ')
                    r_label.font.name = 'Times New Roman'
                    r_label.font.size = Pt(10)
                    r_label.font.italic = True
                    add_seq_field_to_para(cap_p, 'Figur')
                    r_rest = cap_p.add_run('. ' + cap_m.group(3))
                    r_rest.font.name = 'Times New Roman'
                    r_rest.font.size = Pt(10)
                    r_rest.font.italic = True
                else:
                    r = cap_p.add_run(caption)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)
                    r.font.italic = True
            else:
                print(f"  ÅTVARING: Fann ikkje {abs_path}")

    # ── Tabell ──
    elif stripped.startswith('|') and not stripped.startswith('|---'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        i -= 1

        if len(table_lines) >= 3:
            # Parse headers
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            # Skip separator (line 1)
            # Data rows
            data_rows = []
            for tl in table_lines[2:]:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                data_rows.append(cells)

            ncols = len(headers)
            nrows = len(data_rows) + 1

            table = doc.add_table(rows=nrows, cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header (bold, ingen bakgrunnsfarge)
            for j, h in enumerate(headers):
                if j < ncols:
                    cell = table.rows[0].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_cell_content(p, h, font_size=11, bold=True)

            # Data (med inline-matte)
            for ri, row_data in enumerate(data_rows):
                for j, ct in enumerate(row_data):
                    if j < ncols:
                        cell = table.rows[ri + 1].cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        add_cell_content(p, ct, font_size=11)

            set_threeline_borders(table)
            table_count += 1

            # Tabelltittel vert no sett inn FØR tabellen (sjå *Tabell-blokka nedanfor)

    # ── Tabelltittel med SEQ-felt (sett inn OVER tabellen, per kompendiet kap. 3.5.2) ──
    elif stripped.startswith('*Tabell'):
        cap_text = stripped.strip('*').replace('—', '–')
        cap_p = doc.add_paragraph()
        tab_m = re.match(r'(Tabell)\s+(\d+)\.\s*(.*)', cap_text)
        if tab_m:
            r_label = cap_p.add_run('Tabell ')
            r_label.font.name = 'Times New Roman'
            r_label.font.size = Pt(10)
            r_label.font.italic = True
            add_seq_field_to_para(cap_p, 'Tabell')
            r_rest = cap_p.add_run('. ' + tab_m.group(3))
            r_rest.font.name = 'Times New Roman'
            r_rest.font.size = Pt(10)
            r_rest.font.italic = True
        else:
            r = cap_p.add_run(cap_text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.font.italic = True

    # ── Figurtittel (skip — bildetekst er allereie i ![...]) ──
    elif stripped.startswith('*Figur'):
        pass

    # ── Blokkitat ──
    elif stripped.startswith('> '):
        quote = stripped[2:].replace('—', '–')
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Fjern markdown
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', quote)
        r = p.add_run(clean)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        if '**' in quote:
            r.font.bold = True

    # ── Matteblokk (display math) ──
    elif stripped.startswith('$$'):
        math = stripped
        if not stripped.endswith('$$') or stripped == '$$':
            i += 1
            while i < len(lines) and not lines[i].strip().endswith('$$'):
                math += ' ' + lines[i].strip()
                i += 1
            if i < len(lines):
                math += ' ' + lines[i].strip()
        latex = math.replace('$$', '').strip()
        add_math_paragraph(doc, latex)

    # ── HTML (skip) ──
    elif stripped.startswith('<') and stripped.endswith('>'):
        pass

    # ── Vanleg avsnitt ──
    elif stripped and not stripped.startswith('```'):
        p = add_formatted_para(doc, stripped)
        if in_referanseliste:
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0

    i += 1

print(f"  Lagt til {table_count} tabellar, {fig_count} figurar, {math_counter[0]} math-uttrykk")

# ════════════════════════════════════════════════════
# STEG 6: Global formatering
# ════════════════════════════════════════════════════
print("Steg 6: Formatering...")

# Margar
for section in doc.sections:
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# Heading-stilar
for name, size in [('Heading 1', 18), ('Heading 2', 16), ('Heading 3', 14)]:
    if name in doc.styles:
        s = doc.styles[name]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        # Fjern automatisk nummerering frå malen
        pPr = s.element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)

# Normal-stil
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.paragraph_format.line_spacing = 1.5
ns.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── Sidenummerering i footer ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    # PAGE-felt
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)
    for r in [run, run2, run3]:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)

# ════════════════════════════════════════════════════
# STEG 7: Lagre
# ════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"\nFerdig! Lagra som: {OUTPUT}")
print(f"  Totalt avsnitt: {len(doc.paragraphs)}")
print(f"  Tabellar: {len(doc.tables)}")
print("\n  VIKTIG: Opne dokumentet i Word, trykk Ctrl+A -> F9 for å oppdatere")
print("  innholdsfortegnelse, figurliste og tabelliste med sidenummer.")
