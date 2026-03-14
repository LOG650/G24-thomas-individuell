"""
Post-prosessering av pandoc-generert DOCX for HiMolde LOG650-rapport.
Formaterer etter kompendiet og Word-malen:
- Times New Roman 12pt, 1.5 linjeavstand
- Heading 1: 18pt bold, Heading 2: 16pt bold
- Margar: L=3.0, R=2.5, T=2.5, B=2.5 cm
- Tabellar med kantlinjer og tittel under
- Figurar skalert til sidebreidde
- Sideskift før kvart kapittel (Heading 1)
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy
import os, re

INPUT = "LOG650_Rapport_FINAL.docx"
OUTPUT = "LOG650_Rapport_FINAL_formatted.docx"

doc = Document(INPUT)

# ── 1. Margar og sidestorleik ──
for section in doc.sections:
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

# ── 2. Stil-definisjonar ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style_normal.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    style_name = f'Heading {level}'
    if style_name in doc.styles:
        style = doc.styles[style_name]
        style.font.name = 'Times New Roman'
        style.font.bold = True
        style.font.color.rgb = None  # Svart
        if level == 1:
            style.font.size = Pt(18)
            style.paragraph_format.space_before = Pt(24)
            style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            style.font.size = Pt(16)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(8)
        elif level == 3:
            style.font.size = Pt(14)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)

# ── 3. Formater alle avsnitt ──
chapter_keywords = [
    'Kapittel 1', 'Kapittel 2', 'Kapittel 3', 'Kapittel 4',
    'Kapittel 5', 'Kapittel 6', 'Kapittel 7', 'Kapittel 8',
    'Kapittel 9', 'Referanseliste', 'Vedlegg'
]

for para in doc.paragraphs:
    # Sett font på alle runs
    for run in para.runs:
        run.font.name = 'Times New Roman'
        # Sett rFont for å sikre korrekt font i Word
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>')
            rpr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')

    style_name = para.style.name if para.style else 'Normal'

    if style_name == 'Heading 1':
        for run in para.runs:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        # Sideskift før kapitteloverskrifter
        if any(kw in para.text for kw in chapter_keywords):
            ppr = para._element.get_or_add_pPr()
            page_break = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="true"/>')
            # Fjern eksisterande pageBreakBefore
            for existing in ppr.findall(qn('w:pageBreakBefore')):
                ppr.remove(existing)
            ppr.append(page_break)

    elif style_name == 'Heading 2':
        for run in para.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

    elif style_name == 'Heading 3':
        for run in para.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

    else:
        for run in para.runs:
            if run.font.size is None:
                run.font.size = Pt(12)

    # Linjeavstand 1.5 for alt
    para.paragraph_format.line_spacing = 1.5


# ── 4. Formater tabellar ──
def set_table_borders(table):
    """Legg til kantlinjer på tabellen."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Fjern eksisterande borders
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

for table in doc.tables:
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    # rFonts
                    rpr = run._element.get_or_add_rPr()
                    rFonts = rpr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>')
                        rpr.insert(0, rFonts)
                    else:
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.space_before = Pt(2)

            # Skuggelegging på overskriftsrad
            if row_idx == 0:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                tcPr.append(shading)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True


# ── 5. Skaler figurar til sidebreidde ──
# Tilgjengeleg breidde = sidebreidde - venstre marg - høgre marg
avail_width = Cm(21.0) - Cm(3.0) - Cm(2.5)  # = 15.5 cm

body = doc.element.body
for drawing in body.iter(qn('wp:inline')):
    extent = drawing.find(qn('wp:extent'))
    if extent is not None:
        cx = int(extent.get('cx', 0))
        cy = int(extent.get('cy', 0))
        if cx > 0 and cy > 0:
            # Skaler til maks breidde, behold proporsjonar
            max_cx = int(avail_width)
            if cx > max_cx:
                ratio = max_cx / cx
                new_cx = max_cx
                new_cy = int(cy * ratio)
                extent.set('cx', str(new_cx))
                extent.set('cy', str(new_cy))
                # Oppdater også effektExtent om den finst
                for child in drawing:
                    if child.tag.endswith('docPr'):
                        pass
                    # Oppdater graphic-storleik
                graphic = drawing.find(qn('a:graphic'))
                if graphic is not None:
                    graphicData = graphic.find(qn('a:graphicData'))
                    if graphicData is not None:
                        for pic in graphicData.iter():
                            if pic.tag.endswith('}spPr') or pic.tag.endswith('}cNvPr'):
                                pass
                            ext_elem = pic.find(qn('a:ext'))
                            if ext_elem is not None:
                                if ext_elem.get('cx'):
                                    ext_elem.set('cx', str(new_cx))
                                    ext_elem.set('cy', str(new_cy))


# ── 6. Lagre ──
doc.save(OUTPUT)
print(f"Ferdig! Lagra som {OUTPUT}")
print(f"  Avsnitt: {len(doc.paragraphs)}")
print(f"  Tabellar: {len(doc.tables)}")
print(f"  Seksjonar: {len(doc.sections)}")
